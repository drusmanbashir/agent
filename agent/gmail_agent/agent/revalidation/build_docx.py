
from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

from docx import Document

from agent.paths import assert_within_agent_root


_YEAR_DIR_RE = re.compile(r"^\d{4}$")


def infer_active_year(evidence_root: Path) -> int:
    """
    Pick max YYYY folder containing at least one file.
    If none contain files, raise.
    """
    years: List[int] = []
    if not evidence_root.exists():
        raise FileNotFoundError(f"evidence_root not found: {evidence_root}")

    for p in evidence_root.iterdir():
        if p.is_dir() and _YEAR_DIR_RE.match(p.name):
            years.append(int(p.name))

    if not years:
        raise FileNotFoundError(f"no YYYY folders under: {evidence_root}")

    years.sort(reverse=True)
    for y in years:
        yr_dir = evidence_root / str(y)
        if any(x.is_file() for x in yr_dir.rglob("*")):
            return y

    raise FileNotFoundError(f"no evidence files found in any YYYY folder under: {evidence_root}")


class _FormHeadingParser(HTMLParser):
    """
    Minimal HTML parser: extract visible text from headings/legends/labels.
    This gives you a section scaffold for copy/paste into the portal.
    """
    def __init__(self) -> None:
        super().__init__()
        self._capture = False
        self._buf: List[str] = []
        self.headings: List[str] = []
        self._tag_stack: List[str] = []

    def handle_starttag(self, tag, attrs):
        self._tag_stack.append(tag)
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6", "legend", "label"):
            self._capture = True
            self._buf = []

    def handle_endtag(self, tag):
        # pop stack safely
        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6", "legend", "label") and self._capture:
            text = " ".join("".join(self._buf).split()).strip()
            if text and text.lower() not in ("submit", "save", "cancel"):
                self.headings.append(text)
            self._capture = False
            self._buf = []

    def handle_data(self, data):
        if self._capture:
            self._buf.append(data)


def extract_form_sections(form_html: Path) -> List[str]:
    html = form_html.read_text(encoding="utf-8", errors="ignore")
    p = _FormHeadingParser()
    p.feed(html)

    # de-duplicate while preserving order
    seen = set()
    out: List[str] = []
    for h in p.headings:
        k = h.strip()
        if not k:
            continue
        if k in seen:
            continue
        seen.add(k)
        out.append(k)

    # fallback if HTML has no headings
    if not out:
        out = ["Revalidation form sections (could not extract headings)"]
    return out


def list_evidence_files(year_dir: Path) -> List[Path]:
    if not year_dir.exists():
        return []
    return sorted([p for p in year_dir.rglob("*") if p.is_file()])


def guess_bucket(p: Path) -> str:
    s = (p.name + " " + str(p.parent)).lower()
    if any(k in s for k in ("cpd", "certificate", "course", "attendance", "conference", "webinar")):
        return "CPD / Courses"
    if any(k in s for k in ("teach", "lecture", "seminar", "trainer", "faculty")):
        return "Teaching"
    if any(k in s for k in ("audit", "qi", "quality", "governance", "incident", "morbidity", "mortality")):
        return "Quality improvement / Governance"
    if any(k in s for k in ("grant", "fund", "award")):
        return "Grants"
    if any(k in s for k in ("paper", "manuscript", "accepted", "publication", "published", "proof")):
        return "Publications"
    if any(k in s for k in ("review", "reviewer", "referee")):
        return "Peer review"
    return "Other"


def build_revalidation_docx(
    agent_root: Path,
    form_html: Path,
    evidence_root: Path,
    output_dir: Path,
) -> Path:
    year = infer_active_year(evidence_root)
    year_dir = evidence_root / str(year)

    sections = extract_form_sections(form_html)
    files = list_evidence_files(year_dir)

    # bucket evidence
    buckets: dict[str, List[Path]] = {}
    for f in files:
        b = guess_bucket(f)
        buckets.setdefault(b, []).append(f)

    output_dir.mkdir(parents=True, exist_ok=True)
    out_docx = output_dir / f"{year}_revalidation_answers.docx"
    assert_within_agent_root(agent_root, out_docx)

    doc = Document()
    doc.add_heading(f"Revalidation answers pack ({year})", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Evidence folder: {year_dir}")
    doc.add_paragraph(f"Form template: {form_html}")

    doc.add_page_break()
    doc.add_heading("Evidence index (by type guess)", level=2)
    if not files:
        doc.add_paragraph("No files found in the inferred year folder.")
    else:
        for bucket, paths in sorted(buckets.items(), key=lambda x: x[0]):
            doc.add_heading(bucket, level=3)
            for p in paths:
                doc.add_paragraph(str(p.relative_to(evidence_root)), style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Form sections (copy/paste targets)", level=2)
    for sec in sections:
        doc.add_heading(sec, level=3)
        doc.add_paragraph("(paste your final text here)")
        # minimal assist: list possibly relevant evidence under each section title keyword match
        # (kept intentionally crude, no LLM yet)
        hits = []
        key = sec.lower()
        for f in files:
            if key and key in f.name.lower():
                hits.append(f)
        if hits:
            doc.add_paragraph("Potentially related evidence files:", style=None)
            for h in hits[:15]:
                doc.add_paragraph(str(h.relative_to(evidence_root)), style="List Bullet")
        else:
            doc.add_paragraph("MISSING / no obvious evidence linked by filename.")

    doc.save(out_docx)
    return out_docx
